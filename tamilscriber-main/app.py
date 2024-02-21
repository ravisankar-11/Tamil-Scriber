from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify, Response
from werkzeug.utils import secure_filename
import os
from pydub import AudioSegment
import speech_recognition as sr
import moviepy.editor as mp
from googletrans import Translator
from google.transliteration import transliterate_text
import io  
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

app = Flask(__name__)
translator = Translator()


UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
ALLOWED_EXTENSIONS = {'mp3', 'wav', 'mp4', 'avi', 'mkv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
def generate_pdf_from_text(text):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 100, "Generated Subtitles")
    text_lines = text.split('\n')
    y_position = height - 130
    for line in text_lines:
        c.drawString(100, y_position, line)
        y_position -= 15
    c.save()
    buffer.seek(0)
    return buffer.read()

def convert_audio_to_text(audio_path):
    r = sr.Recognizer()

    audio = AudioSegment.from_mp3(audio_path)
    wav_path = os.path.splitext(audio_path)[0] + ".wav"
    audio.export(wav_path, format="wav")

    with sr.AudioFile(wav_path) as source:
        audio = r.record(source)
        try:
            text = r.recognize_google(audio, language='ta-IN')
            return text
        except sr.UnknownValueError:
            return "Unable to transcribe audio"
        except sr.RequestError as e:
            return f"Error: {e}"
def save_text_to_file(text, filename):
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(text)

def generate_subtitles(video_path):
    video = mp.VideoFileClip(video_path)
    audio = video.audio
    audio_path = os.path.join(app.config['UPLOAD_FOLDER'], 'audio.wav')
    audio.write_audiofile(audio_path)

    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_path) as source:
        audio_data = recognizer.record(source)

    try:
        audio_text = recognizer.recognize_google(audio_data, language='ta-IN')
        subtitles = [subtitle.strip() for subtitle in audio_text.split('.')]
        return subtitles
    except sr.UnknownValueError:
        subtitles = ["Speech recognition could not understand the audio."]
        return subtitles
subtitles = []

def generate_docx_from_text(text):
    from docx import Document
    doc = Document()
    doc.add_heading('Generated Subtitles', level=1)
    doc.add_paragraph(text)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream.read()


@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/Rec_Audio.html', methods=['GET', 'POST'])
def Rec_Audio():
    if request.method == 'POST':
        if 'audio' not in request.files:
            return redirect(request.url)
        
        audio_file = request.files['audio']
        
        if audio_file.filename == '':
            return redirect(request.url)
        
        if audio_file:
            audio_path = os.path.join(app.config['UPLOAD_FOLDER'], audio_file.filename)
            audio_file.save(audio_path)
            
            text = convert_audio_to_text(audio_path)
            text_filename = os.path.splitext(audio_file.filename)[0] + ".txt"
            save_text_to_file(text, text_filename)
            
            return render_template('Rec_Audio.html', text=text, text_filename=text_filename)
    
    return render_template('Rec_Audio.html', text=None)

@app.route('/download/<text_filename>', methods=['GET'])
def download(text_filename):
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], text_filename)
    return send_file(output_path, as_attachment=True)
@app.route('/Video.html', methods=['GET', 'POST'])
def Video():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)

        file = request.files['file']

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            video_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(video_path)

            subtitles = generate_subtitles(video_path)

            return render_template('Video.html', video_path=video_path, subtitles=subtitles)

    return render_template('Video.html', video_path=None, subtitles=None)

@app.route('/download_subtitles', methods=['POST'])
def download_subtitles():
    format = request.form.get('format', 'txt')  
    if format == 'txt':
        subtitles_text = "\n".join(subtitles)
        response = Response(subtitles_text, content_type='text/plain')
        response.headers['Content-Disposition'] = 'attachment; filename=subtitles.txt'
        return response
    elif format == 'pdf':
        pdf_content = generate_pdf_from_text("\n".join(subtitles))
        response = Response(pdf_content, content_type='application/pdf')
        response.headers['Content-Disposition'] = 'attachment; filename=subtitles.pdf'
        return response
    elif format == 'docx':
        docx_content = generate_docx_from_text("\n".join(subtitles))
        response = Response(docx_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response.headers['Content-Disposition'] = 'attachment; filename=subtitles.docx'
        return response
@app.route('/live_audio.html', methods=['GET'])
def live_audio_page():
    return render_template('live_audio.html')
@app.route('/start_live_recording', methods=['POST'])
def start_live_recording():
    recognizer = sr.Recognizer()
    microphone = sr.Microphone()

    text = ""
    
    with microphone as source:
        while True:
            audio_data = recognizer.listen(source, timeout=None)
            try:
                result = recognizer.recognize_google(audio_data, language="ta-IN")
                text += result + " "
            except sr.UnknownValueError:
                pass

            if request.form.get('stop'):
                break

    return jsonify({"text": text})

@app.route('/transliterate.html', methods=['GET', 'POST'])
def transliterate():
    source_text = ""
    transliteration_output = ""
    translation_output = ""

    if request.method == "POST":
        source_text = request.form["source_text"]

        translated = translator.translate(source_text, src="en", dest="ta")
        translation_output = translated.text

        transliteration_output = transliterate_text(source_text, lang_code='ta')

    return render_template("transliterate.html", source_text=source_text,
                           transliteration_output=transliteration_output, translation_output=translation_output)

if __name__ == '__main__':
    app.run(debug=True)